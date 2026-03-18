# pipeline/writer.py
from docx import Document
from pathlib import Path

def write_docx(text_blocks, output_path: Path):
    doc = Document()
    for block_type, content in text_blocks:
        if block_type == "table":
            doc.add_paragraph("[Table Detected]")
            doc.add_paragraph(content)
        else:
            doc.add_paragraph(content)
    doc.save(output_path)
